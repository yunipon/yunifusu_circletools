from docx import Document
import os
import sys


def remove_and_reduce_blank_lines(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    buffer = []
    for para in doc.paragraphs:
        text = para.text.strip()

        if text == "":
            buffer.append("")  # 空白行を一時バッファにためる
        else:
            # バッファ処理：空白行が1つだけなら削除、2つ以上なら1つ残す
            if len(buffer) > 1:
                new_doc.add_paragraph("")  # 1行だけ空白を追加
            # len(buffer) == 1 → 追加しない（削除）
            buffer = []  # バッファをクリア

            # 通常行を追加
            new_doc.add_paragraph(text)

    # 最後に空白行で終わっていた場合にも対応（任意で追加しない仕様）
    new_doc.save(output_path)
    print(f"✅ 空白行を整理して保存しました: {output_path}")


# スクリプトの実ファイルのある場所を取得（PyInstallerと通常どちらも対応）
def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# 実行パス取得
base_dir = get_base_dir()
output_dir = os.path.join(base_dir, "【output】")
input_file = os.path.join(base_dir, "00_input.docx")
output_file = os.path.join(output_dir, "output_txt_blankline.docx")

if os.path.exists(input_file):
    remove_and_reduce_blank_lines(input_file, output_file)
else:
    print(f"❌ ファイルが見つかりません: {input_file}")


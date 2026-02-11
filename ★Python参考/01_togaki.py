from docx import Document
import re
import os
import sys

# スクリプトの実ファイルのある場所を取得（PyInstallerと通常どちらも対応）
def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def clean_paragraph_text(text):
    # 半角カッコを全角に変換
    text = text.replace("(", "（").replace(")", "）")

    # ト書きで始まる行パターン
    delete_line_start_patterns = [
        r"^\s*//.*", r"^\s*◇.*", r"^\s*■.*", r"^\s*◆.*", r"^\s*□.*", r"^\s*※.*",
        r"^\s*SE[:：].*"
    ]
    if any(re.match(pattern, text) for pattern in delete_line_start_patterns):
        return " "  # 空白行に置換

    # 括弧内の文字削除
    bracket_patterns = [
        r"【[^】]*】",     # 【…】
        r"（[^）]*）",     # （…）
        r"\([^)]*\)"       # (…) ← 半角括弧も安全のため処理
    ]
    for pattern in bracket_patterns:
        text = re.sub(pattern, '', text)

    # 全角・半角スペースを削除
    cleaned = re.sub(r"[ 　]", '', text).strip()

    return " " if cleaned == "" else cleaned

def remove_leading_spaces(text):
    return re.sub(r"^[ 　]+", "", text)

def process_docx(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        original_text = para.text
        cleaned_text = clean_paragraph_text(original_text)
        cleaned_text = remove_leading_spaces(cleaned_text)
        new_doc.add_paragraph(cleaned_text)

    new_doc.save(output_path)


# 実行パス取得
base_dir = get_base_dir()
input_docx = os.path.join(base_dir, "00_input.docx")
output_dir = os.path.join(base_dir, "【output】")
output_docx = os.path.join(output_dir, "output_cleaned.docx")

# 出力フォルダ作成
os.makedirs(output_dir, exist_ok=True)

# 実行
if os.path.exists(input_docx):
    process_docx(input_docx, output_docx)
    print(f"✅ 完了：'{output_docx}' を保存しました。")
else:
    print(f"❌ エラー：'{input_docx}' が見つかりません。")
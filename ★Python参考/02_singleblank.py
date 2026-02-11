from docx import Document
import os
import sys
import re

def remove_extra_blank_lines(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    previous_was_empty = False

    for para in doc.paragraphs:
        text = para.text.strip()

        if text == "":
            if previous_was_empty:
                continue  # 2つ以上連続する空行はスキップ
            else:
                new_doc.add_paragraph("")  # 1つ目の空行は許可
                previous_was_empty = True
        else:
            new_doc.add_paragraph(text)
            previous_was_empty = False

    new_doc.save(output_path)
    print(f"✅ 空行を整理して保存しました: {output_path}")

# スクリプトの実ファイルのある場所を取得（PyInstallerと通常どちらも対応）
def get_base_dir():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

# 実行パス取得
base_dir = get_base_dir()
output_dir = os.path.join(base_dir, "【output】")
input_file = os.path.join(output_dir, "output_cleaned.docx")
output_file = os.path.join(output_dir, "output_single_blanklines.docx")

if os.path.exists(input_file):
    remove_extra_blank_lines(input_file, output_file)
else:
    print(f"❌ ファイルが見つかりません: {input_file}")

